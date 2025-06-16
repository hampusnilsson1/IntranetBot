import logging
import re
from bs4 import BeautifulSoup
import docx
import pdfplumber
from playwright.sync_api import sync_playwright
from dotenv import load_dotenv
import os

import requests

logging.getLogger("pdfminer").setLevel(logging.ERROR)

UNWANTED_TAGS = [  # Due to site having no main element and inconsistency
    "header",
    "footer",
    "nav",
    "aside",
    "script",
    "style",
    "noscript",
    "iframe",
]
UNWANTED_CLASSES = [
    "tm-header",
    "tm-header-mobile",
    "tm-footer",
    "tm-sidebar",
    "uk-nav",
    "cookie",
    "search",
    "sidebarmenu",
]
UNWANTED_IDS = ["tm-header", "tm-footer", "tm-sidebar", "cookie", "assistant"]

load_dotenv("../data/COOKIE.env")
COOKIE_NAME = os.getenv("COOKIE_NAME")
COOKIE_VALUE = os.getenv("COOKIE_VALUE")


def scrap_site(page_url, cookie_name, cookie_value):
    # Playwright Test
    with sync_playwright() as p:  # "playwright install" Needed in Dockerfile in future
        browser = p.chromium.launch(headless=True)
        context = browser.new_context()

        # Set your cookies here
        context.add_cookies(
            [
                {
                    "name": cookie_name,
                    "value": cookie_value,
                    "domain": "intranet.falkenberg.se",
                    "path": "/",
                    "httpOnly": True,
                    "secure": True,
                }
            ]
        )

        page = context.new_page()
        page.goto(page_url)

        # Wait for the JavaScript to load ms
        page.wait_for_timeout(3000)

        # Check if page was loaded successfully
        current_url = page.url
        if current_url.startswith(
            "https://idp.falkenberg.se/saml/authenticate/samldispatchi"
        ):
            print("Error: Invalid cookie. Redirected to login page.")
            browser.close()
            return None

        # Grab the page content and format it
        content = page.content()
        browser.close()

        soup = BeautifulSoup(content, "html.parser")

        title = soup.title.string if soup.title else "No title found"

        main_content = soup.find("div", id="tm-main")
        if not main_content:
            main_content = soup.find("div", class_="tm-page") or soup.find("body")

        # Filter unnecessary elements
        for tag in UNWANTED_TAGS:
            for match in main_content.find_all(tag):
                match.decompose()
        for class_name in UNWANTED_CLASSES:
            for match in soup.find_all(class_=lambda c: c and class_name in c):
                match.decompose()
        for id in UNWANTED_IDS:
            for match in main_content.find_all(id=id):
                match.decompose()

        if main_content:
            for cookie_div in main_content.find_all(
                "div", id=re.compile("cookie", re.IGNORECASE)
            ):
                cookie_div.decompose()
            texts = " ".join(main_content.stripped_strings)
        else:
            texts = "Main or Page not found or empty."

        results = []
        results.append({"url": page_url, "title": title, "texts": texts})

        # Site Pdfs
        pdf_links = soup.find_all(
            "a",
            href=re.compile(
                r"(^/alla-dokument/|^https://intranet\.falkenberg\.se/alla-dokument/|\.pdf$)",
                re.IGNORECASE,
            ),
        )
        for link in pdf_links:
            pdf_url = link["href"]
            if pdf_url.startswith("/"):
                pdf_url = "https://intranet.falkenberg.se" + pdf_url

            if pdf_url.startswith("https://intranet.falkenberg.se/alla-dokument/"):
                pdf_url = pdf_url + "/file"

            pdf_text = scrap_pdf(pdf_url=pdf_url)
            if pdf_text:
                results.append(
                    {
                        "url": pdf_url,
                        "title": link.text.strip() or "No title",
                        "texts": pdf_text,
                        "source_url": page_url,
                    }
                )
        return results


def scrap_pdf(pdf_url):
    file_path = "temp.pdf"
    try:
        if pdf_url.startswith("https://intranet.falkenberg.se"):
            response = requests.get(
                pdf_url, cookies={COOKIE_NAME: COOKIE_VALUE}, timeout=10
            )
        else:
            response = requests.get(pdf_url)
        response.raise_for_status()

        content_type = response.headers.get("Content-Type", "").lower()
        if (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in content_type
            or "application/msword" in content_type
            or "application/octet-stream" in content_type
        ):
            # Hantera docx-fil
            file_path = "temp.docx"
            with open(file_path, "wb") as f:
                f.write(response.content)

            doc = docx.Document(file_path)
            text_content = []
            for para in doc.paragraphs:
                text_content.append(para.text)

            # Eventuellt ta bort?
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            return " ".join(text_content) if text_content else "No text found in DOCX"
        elif "pdf" in content_type:
            # Hantera pdf-fil
            with open(file_path, "wb") as f:
                f.write(response.content)

            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return " ".join(text_content) if text_content else "No text found in PDF"
        else:
            return
    except (requests.exceptions.RequestException, Exception) as e:
        logging.info(f"Error fetching or processing PDF from {pdf_url}: {str(e)}")
        return "Error fetching or processing PDF"
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
